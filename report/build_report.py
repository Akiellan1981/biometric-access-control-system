# -*- coding: utf-8 -*-
"""Generate the project report .docx for the Biometric Access Control System."""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "Project-Report.docx")
OUT_DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop", "Project-Report.docx")

NAVY = RGBColor(0x1F, 0x38, 0x64)
HDR_BG = "1F3864"
ALT_BG = "EAF1F8"
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- page + base styles ----
sec = doc.sections[0]
sec.page_width = Mm(210); sec.page_height = Mm(297)          # A4
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin = sec.bottom_margin = Mm(18)
CONTENT_MM = 170

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12

for hid, sz in (("Heading 1", 15), ("Heading 2", 12.5)):
    st = doc.styles[hid]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.bold = True
    st.font.color.rgb = NAVY
    st.paragraph_format.space_before = Pt(10); st.paragraph_format.space_after = Pt(4)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_fixed(table):
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)


def add_table(headers, rows, widths_mm, font_pt=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.allow_autofit = False
    set_fixed(t)
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].width = Mm(widths_mm[i]); shade(hc[i], HDR_BG)
        p = hc[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(font_pt); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Mm(widths_mm[i])
            if ri % 2 == 1:
                shade(cells[i], ALT_BG)
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val)); r.font.size = Pt(font_pt)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p


def add_page_number_footer():
    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Dual-Modal Biometric Access Control & Attendance System  |  Page ").font.size = Pt(8)
    run = fp.add_run()
    for typ, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if typ:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), typ); run._r.append(fc)
        else:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = txt; run._r.append(it)
    run.font.size = Pt(8)


# ============================ COVER ============================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(30)
r = title.add_run("Dual-Modal Biometric Access Control\n& Attendance System")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Project Execution Report"); r.font.size = Pt(13); r.font.color.rgb = GREY

line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
line.paragraph_format.space_before = Pt(2)
r = line.add_run("Raspberry Pi 5  ·  Face Recognition + Fingerprint  ·  Anti-Spoofing  ·  Offline  ·  LAN Dashboard")
r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GREY

doc.add_paragraph().paragraph_format.space_after = Pt(8)

meta = add_table(
    ["Field", "Detail"],
    [
        ["Prepared by", "Project Lead (you) – see Section A2 for the team"],
        ["Document version", "1.1"],
        ["Date prepared", "24 June 2026"],
        ["Planned start", "14 July 2026 (after 10 July 2026)"],
        ["Planned completion", "19 September 2026 (about 10 weeks)"],
        ["Deployment", "Indoor, single room; fully offline"],
        ["Software status", "Complete — all modules written, dev-tested on Windows laptop"],
        ["Hardware status", "Pending — Pi 5 deployment after 14 July 2026"],
    ],
    [45, 125], font_pt=10,
)

doc.add_paragraph().add_run().add_break()

# ============================ 1. OVERVIEW ============================
doc.add_heading("1. Project overview (what & why)", level=1)
doc.add_paragraph(
    "What: A standalone device built on a single Raspberry Pi 5 that identifies a person by "
    "either their face or their fingerprint, grants or denies access, records attendance, and "
    "shows all activity on a web dashboard. It works completely offline – no internet or cloud."
)
doc.add_paragraph(
    "Why: To provide one low-cost, private, reliable unit that logs attendance accurately and "
    "blocks fake entries made with printed or phone-screen photos, while keeping all biometric "
    "data on the device."
)
for txt in [
    ("Dual identity – ", "face OR fingerprint is enough to be logged and let in."),
    ("Anti-spoofing – ", "rejects printed and phone-displayed photos."),
    ("Tamper-evident logs – ", "SQLite database with on-demand Excel export."),
    ("Intruder capture – ", "saves a photo of any denied face, tagged unknown vs spoof."),
    ("Voice + dashboard – ", "spoken result through a USB speaker; LAN dashboard behind a login."),
]:
    bullet(txt[1], bold_lead=txt[0])


# ---- Current status note ----
doc.add_heading("Current development status (as of 24 June 2026)", level=2)
doc.add_paragraph(
    "All software has been written and is running on a Windows development laptop before the "
    "Raspberry Pi hardware is procured. The table below shows what is complete and what is "
    "pending for the Pi deployment phase."
)
add_table(
    ["Module / feature", "Status", "Notes"],
    [
        ["Face detection (YuNet)", "Complete", "Tested with laptop webcam, CAP_DSHOW backend on Windows"],
        ["Face recognition (SFace)", "Complete", "128-D cosine-similarity embeddings; gallery reload on enroll"],
        ["Liveness (MiniFASNet + MediaPipe blink)", "Complete (weights pending)", "Code & pipeline ready; ONNX weights downloaded via script; dev_mode disables on laptop"],
        ["Fingerprint driver (mock)", "Complete", "Mock backend tested; real pyfingerprint driver written, needs Pi hardware"],
        ["3-thread pipeline (face + finger + decision)", "Complete", "Producer-consumer queue; shared per-person cooldown (30 s)"],
        ["Encrypted template storage (Fernet / SQLite)", "Complete", "Key auto-created; biometric blobs encrypted at rest"],
        ["Web dashboard (FastAPI + HTMX)", "Complete", "Logs, intruders, people, settings, export, login; auto-refresh every 4 s"],
        ["Live 'Try' page (webcam Register + Run)", "Complete", "Browser-based enroll + real-time recognition + live log"],
        ["Access decision & logging", "Complete", "Tested: grant/deny-unknown/deny-spoof/cooldown all logged correctly"],
        ["Excel attendance export", "Complete", "openpyxl; Name, Date, Time, Method, Result, Score"],
        ["Test suite (pytest)", "Complete", "26 tests green"],
        ["Voice clips (WAV)", "Pending", "Record/source 3 clips (registered.wav / access_granted.wav / access_denied.wav)"],
        ["Pi OS + library install", "Pending", "requirements-pi.txt ready; install after hardware arrives"],
        ["Fingerprint sensor wiring", "Pending", "R307/R503 UART to GPIO; switch driver: mock → pyfingerprint"],
        ["dev_mode → false (real security)", "Pending", "Set config.yaml dev_mode: false on Pi before go-live"],
        ["systemd autostart", "Pending", "One-liner unit file; enables on boot"],
        ["On-site threshold tuning", "Pending", "Calibrate recognition_cosine_thr and live_score_thr under real lighting"],
    ],
    [68, 28, 74], font_pt=9,
)

doc.add_paragraph().add_run().add_break()

# ============================ PART A ============================
pa = doc.add_heading("Part A – Before project execution", level=1)

# A1
doc.add_heading("A1. Project solution document (material list & budget)", level=2)
doc.add_paragraph("Chosen solution in short:")
for txt in [
    ("Compute: ", "Raspberry Pi 5 (8 GB), CPU only – no extra accelerator."),
    ("Face: ", "YuNet (detect) + SFace (recognise) + MiniFASNet & MediaPipe blink (liveness)."),
    ("Fingerprint: ", "R307/R503-class module that matches the print on the sensor itself."),
    ("Data: ", "SQLite database + Excel export; biometric templates encrypted at rest."),
    ("Output: ", "USB speaker for voice; optional relay to drive an electric lock."),
    ("Dashboard: ", "FastAPI with Jinja2 + HTMX, served on the LAN, behind a login."),
]:
    bullet(txt[1], bold_lead=txt[0])
p = doc.add_paragraph(); p.add_run("All software is free and open-source, so the software cost is ₹0.").italic = True

doc.add_paragraph().paragraph_format.space_after = Pt(2)
p = doc.add_paragraph(); r = p.add_run("Material list & budget"); r.bold = True
add_table(
    ["#", "Item", "Qty", "Unit (USD)", "Total (USD)", "Total (₹)", "Purpose"],
    [
        ["1", "Raspberry Pi 5, 8 GB", "1", "80", "80", "6,800", "Main compute unit"],
        ["2", "Active cooler", "1", "5", "5", "425", "Sustained-load cooling"],
        ["3", "Pi Camera Module 3", "1", "25", "25", "2,125", "Face capture"],
        ["4", "Fingerprint module (R307/R503)", "1", "15", "15", "1,275", "Fingerprint, on-sensor match"],
        ["5", "USB-to-TTL adapter", "1", "3", "3", "255", "Sensor wiring (if not GPIO UART)"],
        ["6", "USB mini-speaker", "1", "6", "6", "510", "Voice output"],
        ["7", "27 W USB-C power supply", "1", "12", "12", "1,020", "Power"],
        ["8", "USB-SSD 64 GB", "1", "25", "25", "2,125", "Durable log storage"],
        ["9", "microSD 32 GB", "1", "8", "8", "680", "Boot / OS"],
        ["10", "Opto-isolated relay (optional)", "1", "5", "5", "425", "Electric lock / strike"],
    ],
    [8, 50, 10, 18, 20, 18, 46], font_pt=9,
)
add_table(
    ["Budget summary", "USD", "₹ (approx)"],
    [
        ["Hardware subtotal", "184", "15,640"],
        ["Contingency (15%)", "28", "2,346"],
        ["Software & tools (open-source)", "0", "0"],
        ["Total project budget", "212", "~18,000"],
    ],
    [90, 40, 40], font_pt=9.5,
)
p = doc.add_paragraph()
r = p.add_run("Note: prices are approximate; verify current local prices. ₹ shown at about ₹85 per USD.")
r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY

# A2
doc.add_heading("A2. Project members & responsibilities", level=2)
doc.add_paragraph("The work is shared across three members. The Project Lead (you) owns overall delivery.")
add_table(
    ["Member / role", "Name", "Main responsibilities"],
    [
        ["M1 – Project Lead & Software Integration", "(You)",
         "Overall ownership and decisions; software architecture; decision/logging engine; "
         "dashboard; final integration; on-site tuning; reporting."],
        ["M2 – Hardware & Electronics", "____________",
         "Procurement; Pi assembly and cooling; wiring of camera, fingerprint sensor, speaker "
         "and relay; power and thermal checks; systemd autostart; field installation."],
        ["M3 – Vision/ML & Testing", "____________",
         "Face pipeline (YuNet/SFace); liveness (MiniFASNet + MediaPipe blink); enrollment "
         "quality; threshold calibration; running the test plan and recording results."],
    ],
    [55, 30, 85], font_pt=9.5,
)

# A3
doc.add_heading("A3. Project timeline", level=2)
doc.add_paragraph(
    "Work begins on 14 July 2026 (after 10 July 2026) and is planned to finish on 19 September "
    "2026 – about ten weeks. Phases run one after another."
)
add_table(
    ["Phase", "Dates (2026)", "Deliverable", "Status"],
    [
        ["PRE – Software development", "Jun 2026 (done)", "All modules written & dev-tested on Windows laptop; 26 pytest green", "DONE"],
        ["0 – Procurement & Pi setup", "14 Jul – 25 Jul", "Hardware bought; Pi OS, Python, libraries, code deployed to Pi", "Pending"],
        ["1 – Core recognition on Pi", "28 Jul – 08 Aug", "Camera + YuNet/SFace running on Pi; SQLite verified on USB-SSD", "Pending"],
        ["2 – Liveness on Pi", "11 Aug – 16 Aug", "MiniFASNet weights loaded on Pi; blink + score calibrated", "Pending"],
        ["3 – Fingerprint & decision", "18 Aug – 23 Aug", "Real fingerprint sensor wired; driver switched to pyfingerprint", "Pending"],
        ["4 – Intruder, voice & deploy mode", "25 Aug – 30 Aug", "Voice clips recorded; dev_mode: false; relay wired", "Pending"],
        ["5 – Dashboard LAN test", "01 Sep – 12 Sep", "Dashboard verified from another PC; login + export working", "Pending"],
        ["6 – Tune, test, go-live", "15 Sep – 19 Sep", "On-site threshold tuning; soak test; systemd autostart; handover", "Pending"],
    ],
    [42, 32, 76, 20], font_pt=9,
)

doc.add_paragraph().add_run().add_break()

# ============================ PART B ============================
doc.add_heading("Part B – After project execution", level=1)

# B1
doc.add_heading("B1. Project execution process (what)", level=2)
doc.add_paragraph(
    "What it is: the ordered build that turns parts and code into one working unit. Steps:"
)
for txt in [
    "Set up the Raspberry Pi: OS, Python 3.11, and all libraries; create the code repository with a single config file for every tunable.",
    "Build module by module following the phases (recognition, liveness, fingerprint, logging, intruder, voice, dashboard); commit each step.",
    "Wire the three working threads – face, fingerprint, and one decision/logging thread – around a single shared queue so database writes stay safe.",
    "Enrol test users: a few face samples plus one fingerprint per person, saved into one unified record.",
    "Connect the dashboard, the USB-speaker voice, and (if used) the relay.",
    "Keep every access decision on the device – no internet at any point.",
]:
    numbered(txt)

# B2
doc.add_heading("B2. Trial run & verify performance (how)", level=2)
doc.add_paragraph(
    "How we check it: run the unit in the real room and test each requirement with a simple "
    "pass/fail. Failures are fixed and re-tested; thresholds are tuned on-site."
)
add_table(
    ["Test", "How to check", "Pass criteria"],
    [
        ["Recognition accuracy", "Enrol 5–10 people; each tries ~20 times in real lighting", "Correct person let in > 95%; no wrong person let in"],
        ["Anti-spoofing", "Show a printed photo and a phone-screen photo of an enrolled person", "Both rejected and tagged 'denied-spoof'; image saved"],
        ["Dual modality", "Same person uses face, then fingerprint", "Both grant and log to the same person record"],
        ["No duplicates", "Trigger face and finger within 30 seconds", "Logged only once (cooldown works)"],
        ["Speed", "Time from face seen to decision", "Under about 1 second"],
        ["Intruder capture", "An unknown person stands in front", "Photo saved to 'unauthorized/', tagged 'denied-unknown'"],
        ["Voice", "Each outcome (registered / granted / denied)", "Correct clip plays on the USB speaker"],
        ["Dashboard", "Open from another PC on the same network", "Login works; logs, gallery, people, export all work"],
        ["Thermal & stability", "Run continuously for a few hours", "Chip temperature stays below throttling; no crash"],
    ],
    [38, 70, 62], font_pt=9,
)

# B3
doc.add_heading("B3. Final implementation (how)", level=2)
doc.add_paragraph("How we go live once trials pass:")
for txt in [
    "Mount the unit so light falls on faces, never with the camera facing a window.",
    "Use the USB-SSD for logs; switch the fingerprint driver from 'mock' to the real 'pyfingerprint' sensor.",
    "Set the dashboard login password and bind the dashboard to the LAN only.",
    "If a lock is used, wire the relay on its own power rail and set its active level and hold time.",
    "Enable systemd autostart so the system runs automatically on boot.",
    "Enrol all real users and run on-site threshold tuning under the actual lighting.",
    "Optionally turn on 'fingerprint required for unlock' for security-critical doors.",
    "Hand over the unit with a short operator guide.",
]:
    numbered(txt)

# B4
doc.add_heading("B4. Project completion document (what / why / how)", level=2)
p = doc.add_paragraph(); r = p.add_run("What was delivered: "); r.bold = True
p.add_run(
    "a working dual-modal access and attendance unit, a LAN dashboard, and documentation – "
    "meeting the project's acceptance criteria."
)
p = doc.add_paragraph(); r = p.add_run("Why it is complete: "); r.bold = True
p.add_run(
    "all trial tests in Section B2 pass under deployment lighting; the system runs offline, "
    "logs correctly, rejects printed and phone-photo spoofs, and stays stable during a soak test."
)
p = doc.add_paragraph(); r = p.add_run("How to operate and maintain it:"); r.bold = True
for txt in [
    ("Daily use: ", "it starts on boot; staff simply present their face or finger."),
    ("Add / remove people: ", "use the dashboard 'People' page (Register mode)."),
    ("Get attendance: ", "dashboard 'Export' button downloads an Excel file."),
    ("Review intruders: ", "dashboard 'Intruders' gallery; spoof attempts are flagged."),
    ("Maintenance: ", "keep SSD space free, review logs, re-tune thresholds if lighting changes; "
     "old intruder images auto-delete after the retention limit."),
]:
    bullet(txt[1], bold_lead=txt[0])

p = doc.add_paragraph(); r = p.add_run("Handover items: "); r.bold = True
p.add_run("device; source code and config; this report; the operator guide; the admin password.")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
p = doc.add_paragraph(); r = p.add_run("Sign-off"); r.bold = True
add_table(
    ["Role", "Name", "Signature", "Date"],
    [
        ["Project Lead (M1)", "(You)", "", ""],
        ["Hardware & Electronics (M2)", "", "", ""],
        ["Vision/ML & Testing (M3)", "", "", ""],
    ],
    [55, 50, 40, 25], font_pt=9.5,
)

add_page_number_footer()
doc.save(OUT)
print("SAVED:", OUT)
try:
    doc.save(OUT_DESKTOP)
    print("SAVED:", OUT_DESKTOP)
except PermissionError:
    print("NOTE: Desktop file is open in Word — close it and re-run, or use the copy at:", OUT)
